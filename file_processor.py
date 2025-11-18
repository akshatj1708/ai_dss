from pathlib import Path

import pandas as pd
import PyPDF2
from docx import Document

from PIL import Image
import pytesseract

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.xls', '.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
        self.data = None
        self.file_info = {}
    
    def process_file(self, file_path):   #MAIN FILE PROCESSOR FUNCTION
        """Main file processing function"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported format: {extension}")
        
        # Store file info
        self.file_info = {
            'name': file_path.name,
            'size': file_path.stat().st_size,
            'extension': extension
        }
        
        # Process based on file type
        if extension in ['.csv', '.xlsx', '.xls']:
            return self._process_dataframe(file_path)
        elif extension in ['.pdf', '.docx', '.doc', '.txt']:
            return self._process_text(file_path)
        elif extension in ['.png', '.jpg', '.jpeg']:
            return self._process_image(file_path)
        
        # If the extension is in supported_formats but no specific processing method is found
        raise ValueError(f"No processing method found for file type: {extension}")
    
    def _process_dataframe(self, file_path):
        """Process dataframe files (CSV, Excel)"""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                sheet_info = None
            else:  # Excel files
                excel_file = pd.ExcelFile(file_path)
                sheet_names = excel_file.sheet_names
                df = pd.read_excel(file_path, sheet_name=sheet_names[0])
                sheet_info = {
                    'sheet_name': sheet_names[0],
                    'available_sheets': sheet_names if len(sheet_names) > 1 else None
                }
            
            self.data = df
            result = {
                'type': 'dataframe',
                'data': df,
                'shape': df.shape,
                'columns': list(df.columns),#conversion to list bcz return type of df.columns is pandas.Index
                'summary': df.describe().to_dict()
            }
            
            if sheet_info:
                result.update(sheet_info)
                if sheet_info['available_sheets']:
                    result['message'] = f"Multiple sheets found. Using first sheet '{sheet_names[0]}'. Available sheets: {', '.join(sheet_names)}"
            
            return result
            
        except Exception as e:
            raise Exception(f"Dataframe processing failed: {str(e)}")
    
    def _process_text(self, file_path):
        """Process text-based files (PDF, DOCX, TXT)"""
        try:
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                metadata = {'pages': len(reader.pages)}
                
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                metadata = {'paragraphs': len(doc.paragraphs)}
                
            else:  # TXT files
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                metadata = {'lines': len(text.split('\n'))}
            
            return {
                'type': 'text',
                'data': text,
                'word_count': len(text.split()),
                **metadata  #** is a dictionary unpacking operatorIt takes all key-value pairs
            }               #from metadata dict which varies based on file-type and adds them directly into the
                            #returned dictionary    
        except Exception as e:
            raise Exception(f"Text processing failed: {str(e)}")
    
    def _process_image(self, file_path):
        """Process image files with OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                'type': 'text',
                'data': text,
                'image_size': image.size,
                'word_count': len(text.split())
            }
        except Exception as e:
            raise Exception(f"Image processing failed: {str(e)}")